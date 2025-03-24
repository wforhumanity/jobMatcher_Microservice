import io
import docx
from typing import Optional

def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a .docx file.
    
    Args:
        file_content: The binary content of the .docx file
        
    Returns:
        The extracted text
    """
    doc = docx.Document(io.BytesIO(file_content))
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return '\n'.join(full_text)

def process_resume_file(file_content: bytes, filename: str) -> Optional[str]:
    """
    Process a resume file and extract its text content.
    
    Args:
        file_content: The binary content of the file
        filename: The name of the file
        
    Returns:
        The extracted text or None if the file format is not supported
    """
    if filename.lower().endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif filename.lower().endswith('.txt'):
        return file_content.decode('utf-8')
    else:
        return None
